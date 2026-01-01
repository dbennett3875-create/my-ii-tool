"""
I&I Analysis Tool V2 - Enhanced Edition
Complete feature set with advanced analysis capabilities
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime
from scipy import signal
from sklearn.ensemble import IsolationForest
import io
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

# Page setup
st.set_page_config(page_title="I&I Analysis Tool", page_icon="💧", layout="wide")

st.title("💧 I&I Analysis Tool - Enhanced")
st.markdown("**Infiltration & Inflow Analysis for Wastewater Systems**")

# Sidebar
with st.sidebar:
    st.header("⚙️ Settings")
    rain_threshold = st.slider("Rain Event Threshold (in)", 0.01, 0.5, 0.1, 0.01)
    baseline_percentile = st.slider("Baseline Percentile", 5, 25, 10)

    st.markdown("---")
    st.subheader("🔬 Advanced Analysis")
    enable_hydrograph_sep = st.checkbox("Hydrograph Separation", value=True)
    enable_rdii = st.checkbox("RDII Coefficients", value=True)
    enable_anomaly = st.checkbox("Anomaly Detection", value=True)

    if enable_anomaly:
        anomaly_method = st.selectbox("Anomaly Method", ["IQR", "Z-Score", "Isolation Forest"])

    if enable_hydrograph_sep:
        alpha = st.slider("Filter Parameter (α)", 0.90, 0.98, 0.925, 0.005,
                         help="Higher values = slower baseflow response")

# Helper functions
def hydrograph_separation(flow_series, alpha=0.925):
    """Recursive digital filter for baseflow separation"""
    flow = flow_series.values
    baseflow = np.zeros(len(flow))
    baseflow[0] = flow[0]

    for i in range(1, len(flow)):
        baseflow[i] = (alpha * baseflow[i-1]) + ((1 - alpha) / 2) * (flow[i] + flow[i-1])
        baseflow[i] = min(baseflow[i], flow[i])

    quickflow = flow - baseflow
    return baseflow, quickflow

def calculate_rdii(rain_events, flow_df, rain_df, baseline, asset_df=None):
    """Calculate RDII coefficients"""
    if not rain_events or asset_df is None or 'area_acres' not in asset_df.columns:
        return None

    total_rain_volume = sum(e['total_rain'] for e in rain_events)
    total_ii_volume = ((flow_df['flow_rate_gpm'] - baseline).clip(lower=0).sum() * 60) / 1000000  # MGal

    # Get total drainage area
    total_area = asset_df['area_acres'].sum() if 'area_acres' in asset_df.columns else 1

    # R-value: ratio of I&I volume to rainfall volume
    rain_volume_mg = total_rain_volume * total_area * 27154 / 1000000  # Convert acre-inches to MGal
    r_value = total_ii_volume / rain_volume_mg if rain_volume_mg > 0 else 0

    # Unit I&I: gallons per inch of rain per acre
    unit_ii = (total_ii_volume * 1000000) / (total_rain_volume * total_area) if (total_rain_volume * total_area) > 0 else 0

    return {
        'r_value': r_value,
        'unit_ii_gal_per_inch_acre': unit_ii,
        'total_ii_volume_mg': total_ii_volume,
        'total_rain_volume_in': total_rain_volume,
        'drainage_area_acres': total_area
    }

def detect_anomalies(flow_series, method='IQR'):
    """Detect flow anomalies"""
    if method == 'IQR':
        Q1 = flow_series.quantile(0.25)
        Q3 = flow_series.quantile(0.75)
        IQR = Q3 - Q1
        lower = Q1 - 1.5 * IQR
        upper = Q3 + 1.5 * IQR
        anomalies = (flow_series < lower) | (flow_series > upper)

    elif method == 'Z-Score':
        z_scores = np.abs((flow_series - flow_series.mean()) / flow_series.std())
        anomalies = z_scores > 3

    else:  # Isolation Forest
        iso_forest = IsolationForest(contamination=0.1, random_state=42)
        anomalies = iso_forest.fit_predict(flow_series.values.reshape(-1, 1)) == -1

    return anomalies

def generate_client_report(rain_events, baseline, flow_df, rain_df, asset_df=None, rdii_results=None,
                          enable_anomaly=False, basin_priority_df=None):
    """Generate a professional Word document report for client presentation"""

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    title = doc.add_heading('Infiltration & Inflow Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.add_run('Wastewater System Performance Evaluation').bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.add_run(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', 1)

    total_rainfall = sum(e['total_rain'] for e in rain_events) if rain_events else 0
    total_ii = sum(e.get('ii_volume_mg', 0) for e in rain_events) if rain_events else 0

    summary_text = f"""
This report presents the results of an Infiltration and Inflow (I&I) analysis conducted on the wastewater collection system.
The analysis period encompasses {len(flow_df)} flow measurements and identified {len(rain_events)} distinct rain events
totaling {total_rainfall:.2f} inches of rainfall.

Key Findings:
• Baseline Dry Weather Flow: {baseline:.1f} GPM
• Number of Rain Events Analyzed: {len(rain_events)}
• Total I&I Volume: {total_ii:.2f} Million Gallons
• Peak Flow Observed: {flow_df['flow_rate_gpm'].max():.1f} GPM
"""

    if rdii_results:
        summary_text += f"""• R-Value (RDII Coefficient): {rdii_results['r_value']:.3f}
• Unit I&I: {rdii_results['unit_ii_gal_per_inch_acre']:.1f} gallons per inch per acre
"""

    doc.add_paragraph(summary_text)

    # Analysis Period
    doc.add_heading('Analysis Period & Data Summary', 2)

    period_start = flow_df['timestamp'].min().strftime("%B %d, %Y")
    period_end = flow_df['timestamp'].max().strftime("%B %d, %Y")

    doc.add_paragraph(f"Analysis Period: {period_start} to {period_end}")
    doc.add_paragraph(f"Total Data Points: {len(flow_df)}")

    if 'basin_id' in flow_df.columns:
        basins = flow_df['basin_id'].unique()
        doc.add_paragraph(f"Number of Basins Analyzed: {len(basins)}")
        doc.add_paragraph(f"Basins: {', '.join(basins)}")

    doc.add_page_break()

    # Key Metrics
    doc.add_heading('Key Performance Metrics', 1)

    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.style = 'Light Grid Accent 1'

    hdr_cells = metrics_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'

    metrics = [
        ('Baseline Dry Weather Flow', f'{baseline:.1f} GPM'),
        ('Peak Flow Rate', f'{flow_df["flow_rate_gpm"].max():.1f} GPM'),
        ('Average Flow Rate', f'{flow_df["flow_rate_gpm"].mean():.1f} GPM'),
        ('Total Rainfall', f'{total_rainfall:.2f} inches'),
        ('Total I&I Volume', f'{total_ii:.2f} Million Gallons'),
        ('Number of Rain Events', str(len(rain_events))),
    ]

    if rdii_results:
        metrics.extend([
            ('R-Value', f'{rdii_results["r_value"]:.3f}'),
            ('Unit I&I', f'{rdii_results["unit_ii_gal_per_inch_acre"]:.1f} gal/in/acre'),
            ('Drainage Area', f'{rdii_results["drainage_area_acres"]:.1f} acres')
        ])

    if enable_anomaly:
        anomaly_count = flow_df['is_anomaly'].sum() if 'is_anomaly' in flow_df.columns else 0
        anomaly_pct = (anomaly_count / len(flow_df)) * 100 if len(flow_df) > 0 else 0
        metrics.append(('Anomalous Readings', f'{anomaly_count} ({anomaly_pct:.1f}%)'))

    for metric_name, metric_value in metrics:
        row_cells = metrics_table.add_row().cells
        row_cells[0].text = metric_name
        row_cells[1].text = metric_value

    doc.add_page_break()

    # Rain Event Details
    if rain_events:
        doc.add_heading('Rain Event Analysis', 1)

        doc.add_paragraph(f"""
A total of {len(rain_events)} rain events were identified during the analysis period. Each event was analyzed to determine
the I&I response characteristics. The table below summarizes key metrics for each event.
        """)

        # Event table
        event_table = doc.add_table(rows=1, cols=6)
        event_table.style = 'Light Grid Accent 1'

        hdr_cells = event_table.rows[0].cells
        hdr_cells[0].text = 'Event'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Rainfall (in)'
        hdr_cells[3].text = 'Peak Flow (GPM)'
        hdr_cells[4].text = 'I&I Volume (MG)'
        hdr_cells[5].text = 'Peak/Baseline'

        for event in rain_events:
            row_cells = event_table.add_row().cells
            row_cells[0].text = str(event.get('event_num', ''))
            row_cells[1].text = event['start_time'].strftime("%m/%d/%Y")
            row_cells[2].text = f"{event['total_rain']:.2f}"
            row_cells[3].text = f"{event.get('peak_flow_gpm', 0):.1f}"
            row_cells[4].text = f"{event.get('ii_volume_mg', 0):.2f}"
            row_cells[5].text = f"{event.get('peak_baseline_ratio', 0):.1f}"

        doc.add_page_break()

    # Basin Prioritization
    if basin_priority_df is not None and not basin_priority_df.empty:
        doc.add_heading('Basin Prioritization', 1)

        doc.add_paragraph("""
Based on I&I volume, infrastructure age, and pipe length, the following basins have been prioritized for remediation efforts:
        """)

        priority_table = doc.add_table(rows=1, cols=4)
        priority_table.style = 'Light Grid Accent 1'

        hdr_cells = priority_table.rows[0].cells
        hdr_cells[0].text = 'Rank'
        hdr_cells[1].text = 'Basin'
        hdr_cells[2].text = 'I&I Volume (MG)'
        hdr_cells[3].text = 'Priority Score'

        for idx, row in basin_priority_df.iterrows():
            row_cells = priority_table.add_row().cells
            row_cells[0].text = str(row.get('Priority Rank', ''))
            row_cells[1].text = str(row['Basin'])
            row_cells[2].text = f"{row['I&I Volume (MG)']:.2f}"
            row_cells[3].text = f"{row.get('Priority Score', 0):.1f}"

        doc.add_page_break()

    # Recommendations Section (Template for Engineer to Edit)
    doc.add_heading('Recommendations', 1)

    doc.add_paragraph("""
[ENGINEER: Please customize this section based on the specific findings for this system]

Based on the analysis results, the following recommendations are provided:
    """)

    # Add numbered recommendations template
    doc.add_paragraph('1. Immediate Actions', style='List Number')
    doc.add_paragraph('[Describe urgent issues requiring immediate attention, e.g., high I&I basins, critical infrastructure concerns]')

    doc.add_paragraph('2. Short-term Improvements (0-12 months)', style='List Number')
    doc.add_paragraph('[List recommended repairs, inspections, or pilot programs to address identified I&I sources]')

    doc.add_paragraph('3. Long-term Planning (1-5 years)', style='List Number')
    doc.add_paragraph('[Outline systematic rehabilitation plan, monitoring programs, or capital improvement projects]')

    doc.add_paragraph('4. Cost-Benefit Considerations', style='List Number')
    doc.add_paragraph('[Discuss economic analysis, ROI projections, or funding opportunities for recommended improvements]')

    doc.add_page_break()

    # Technical Methodology
    doc.add_heading('Technical Methodology', 1)

    methodology_text = """
This analysis employed industry-standard techniques for I&I quantification and characterization:

• Baseline Flow Calculation: Statistical analysis using the lower percentile of dry weather flows
• Rain Event Detection: Automated identification based on rainfall intensity thresholds
• I&I Volume Quantification: Integration of flow above baseline during wet weather periods
• Hydrograph Separation: Recursive digital filter method to distinguish baseflow from quickflow
• RDII Analysis: Calculation of rainfall-dependent infiltration and inflow coefficients
• Anomaly Detection: Statistical methods to identify unusual flow patterns requiring investigation

All calculations were performed using validated algorithms consistent with EPA and industry best practices.
    """

    doc.add_paragraph(methodology_text)

    doc.add_page_break()

    # Conclusion
    doc.add_heading('Conclusion', 1)

    conclusion_text = f"""
[ENGINEER: Please customize this conclusion section]

This comprehensive I&I analysis has quantified the extent of infiltration and inflow within the wastewater collection system.
The identified I&I volume of {total_ii:.2f} million gallons during {len(rain_events)} rain events represents
[INSERT PERCENTAGE]% of the total flow during wet weather conditions.

The findings presented in this report provide a foundation for prioritizing system improvements and optimizing
the performance of the wastewater collection and treatment infrastructure.

For questions regarding this analysis or to discuss the recommendations in detail, please contact:

[INSERT ENGINEERING FIRM NAME]
[INSERT CONTACT INFORMATION]
    """

    doc.add_paragraph(conclusion_text)

    # Save to bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer

# File uploaders
st.header("📤 Upload Data")

col1, col2, col3 = st.columns(3)

with col1:
    flow_file = st.file_uploader("Flow Data (CSV)", type=['csv'], key='flow')

with col2:
    rain_file = st.file_uploader("Rainfall Data (CSV)", type=['csv'], key='rain')

with col3:
    asset_file = st.file_uploader("Asset Data (CSV, optional)", type=['csv'], key='asset')

# Process data if files uploaded
if flow_file and rain_file:
    # Load data
    flow_df = pd.read_csv(flow_file)
    rain_df = pd.read_csv(rain_file)
    asset_df = pd.read_csv(asset_file) if asset_file else None

    # Parse timestamps
    flow_df['timestamp'] = pd.to_datetime(flow_df['timestamp'])
    rain_df['timestamp'] = pd.to_datetime(rain_df['timestamp'])

    st.success(f"✅ Loaded {len(flow_df)} flow records and {len(rain_df)} rainfall records")

    # Analysis button
    if st.button("🔍 Run Analysis", type="primary"):
        with st.spinner("Analyzing..."):

            # Calculate baseline
            baseline = flow_df['flow_rate_gpm'].quantile(baseline_percentile / 100)

            # Hydrograph separation
            if enable_hydrograph_sep:
                if 'basin_id' in flow_df.columns:
                    flow_df['baseflow'] = 0
                    flow_df['quickflow'] = 0
                    for basin in flow_df['basin_id'].unique():
                        mask = flow_df['basin_id'] == basin
                        bf, qf = hydrograph_separation(flow_df.loc[mask, 'flow_rate_gpm'], alpha)
                        flow_df.loc[mask, 'baseflow'] = bf
                        flow_df.loc[mask, 'quickflow'] = qf
                else:
                    baseflow, quickflow = hydrograph_separation(flow_df['flow_rate_gpm'], alpha)
                    flow_df['baseflow'] = baseflow
                    flow_df['quickflow'] = quickflow

            # Anomaly detection
            if enable_anomaly:
                if 'basin_id' in flow_df.columns:
                    flow_df['is_anomaly'] = False
                    for basin in flow_df['basin_id'].unique():
                        mask = flow_df['basin_id'] == basin
                        flow_df.loc[mask, 'is_anomaly'] = detect_anomalies(
                            flow_df.loc[mask, 'flow_rate_gpm'], anomaly_method
                        )
                else:
                    flow_df['is_anomaly'] = detect_anomalies(flow_df['flow_rate_gpm'], anomaly_method)

            # Detect rain events
            rain_events = []
            in_event = False
            event_start = None

            for idx, row in rain_df.iterrows():
                if row['rainfall_in'] >= rain_threshold:
                    if not in_event:
                        event_start = idx
                        in_event = True
                elif in_event:
                    event_end = idx
                    event_data = {
                        'event_num': len(rain_events) + 1,
                        'start_idx': event_start,
                        'end_idx': event_end,
                        'start_time': rain_df.loc[event_start, 'timestamp'],
                        'end_time': rain_df.loc[event_end, 'timestamp'],
                        'total_rain': rain_df.loc[event_start:event_end, 'rainfall_in'].sum(),
                        'peak_rain': rain_df.loc[event_start:event_end, 'rainfall_in'].max()
                    }

                    # Calculate I&I for this event
                    event_flow = flow_df[
                        (flow_df['timestamp'] >= event_data['start_time']) &
                        (flow_df['timestamp'] <= event_data['end_time'])
                    ]
                    if len(event_flow) > 0:
                        event_data['peak_flow_gpm'] = event_flow['flow_rate_gpm'].max()
                        event_data['ii_volume_mg'] = ((event_flow['flow_rate_gpm'] - baseline).clip(lower=0).sum() * 60) / 1000000
                        event_data['peak_baseline_ratio'] = event_data['peak_flow_gpm'] / baseline if baseline > 0 else 0

                    rain_events.append(event_data)
                    in_event = False

            # Calculate RDII if enabled
            rdii_results = None
            if enable_rdii and rain_events:
                rdii_results = calculate_rdii(rain_events, flow_df, rain_df, baseline, asset_df)

            # === RESULTS SECTION ===
            st.markdown("---")
            st.header("📊 Analysis Results")

            # Key metrics
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Baseline Flow", f"{baseline:.1f} GPM")
            col2.metric("Rain Events", len(rain_events))
            if rain_events:
                total_rainfall = sum(e['total_rain'] for e in rain_events)
                col3.metric("Total Rainfall", f"{total_rainfall:.2f} in")
                total_ii = sum(e.get('ii_volume_mg', 0) for e in rain_events)
                col4.metric("Total I&I Volume", f"{total_ii:.2f} MG")

            # RDII Metrics
            if rdii_results:
                st.subheader("📐 RDII Coefficients")
                col1, col2, col3 = st.columns(3)
                col1.metric("R-Value", f"{rdii_results['r_value']:.3f}")
                col2.metric("Unit I&I", f"{rdii_results['unit_ii_gal_per_inch_acre']:.1f} gal/in/acre")
                col3.metric("Drainage Area", f"{rdii_results['drainage_area_acres']:.1f} acres")

            # Anomaly summary
            if enable_anomaly:
                anomaly_count = flow_df['is_anomaly'].sum()
                anomaly_pct = (anomaly_count / len(flow_df)) * 100
                st.info(f"🔍 Detected {anomaly_count} anomalous flow readings ({anomaly_pct:.1f}% of data)")

            # === VISUALIZATIONS ===
            st.subheader("📈 Flow Analysis")

            # Main hydrograph
            fig = make_subplots(
                rows=2, cols=1,
                row_heights=[0.3, 0.7],
                subplot_titles=('Rainfall', 'Flow Rate with Baseflow Separation' if enable_hydrograph_sep else 'Flow Rate'),
                vertical_spacing=0.1
            )

            # Rainfall bars
            fig.add_trace(
                go.Bar(x=rain_df['timestamp'], y=rain_df['rainfall_in'],
                      name='Rainfall', marker_color='steelblue'),
                row=1, col=1
            )

            # Flow lines by basin or total
            if 'basin_id' in flow_df.columns:
                for basin in flow_df['basin_id'].unique():
                    basin_data = flow_df[flow_df['basin_id'] == basin]

                    # Main flow
                    fig.add_trace(
                        go.Scatter(x=basin_data['timestamp'], y=basin_data['flow_rate_gpm'],
                                  mode='lines', name=f'{basin} - Total Flow',
                                  line=dict(width=2)),
                        row=2, col=1
                    )

                    # Baseflow if enabled
                    if enable_hydrograph_sep:
                        fig.add_trace(
                            go.Scatter(x=basin_data['timestamp'], y=basin_data['baseflow'],
                                      mode='lines', name=f'{basin} - Baseflow',
                                      line=dict(dash='dash', width=1.5)),
                            row=2, col=1
                        )

                    # Anomalies if enabled
                    if enable_anomaly:
                        anomaly_data = basin_data[basin_data['is_anomaly']]
                        if len(anomaly_data) > 0:
                            fig.add_trace(
                                go.Scatter(x=anomaly_data['timestamp'], y=anomaly_data['flow_rate_gpm'],
                                          mode='markers', name=f'{basin} - Anomalies',
                                          marker=dict(color='red', size=8, symbol='x')),
                                row=2, col=1
                            )
            else:
                # Single flow trace
                fig.add_trace(
                    go.Scatter(x=flow_df['timestamp'], y=flow_df['flow_rate_gpm'],
                              mode='lines', name='Total Flow', line=dict(width=2, color='darkblue')),
                    row=2, col=1
                )

                if enable_hydrograph_sep:
                    fig.add_trace(
                        go.Scatter(x=flow_df['timestamp'], y=flow_df['baseflow'],
                                  mode='lines', name='Baseflow',
                                  line=dict(dash='dash', width=1.5, color='green')),
                        row=2, col=1
                    )
                    fig.add_trace(
                        go.Scatter(x=flow_df['timestamp'], y=flow_df['quickflow'],
                                  mode='lines', name='Quickflow',
                                  line=dict(dash='dot', width=1, color='orange'),
                                  fill='tonexty'),
                        row=2, col=1
                    )

                if enable_anomaly:
                    anomaly_data = flow_df[flow_df['is_anomaly']]
                    if len(anomaly_data) > 0:
                        fig.add_trace(
                            go.Scatter(x=anomaly_data['timestamp'], y=anomaly_data['flow_rate_gpm'],
                                      mode='markers', name='Anomalies',
                                      marker=dict(color='red', size=8, symbol='x')),
                            row=2, col=1
                        )

            # Baseline line
            fig.add_trace(
                go.Scatter(x=flow_df['timestamp'],
                          y=[baseline] * len(flow_df),
                          mode='lines', name='Statistical Baseline',
                          line=dict(dash='dash', color='red', width=2)),
                row=2, col=1
            )

            # Highlight rain events
            for event in rain_events:
                fig.add_vrect(
                    x0=event['start_time'], x1=event['end_time'],
                    fillcolor="lightblue", opacity=0.2,
                    layer="below", line_width=0,
                    row=2, col=1
                )

            fig.update_xaxes(title_text="Time", row=2, col=1)
            fig.update_yaxes(title_text="Rainfall (in)", row=1, col=1)
            fig.update_yaxes(title_text="Flow Rate (GPM)", row=2, col=1)

            fig.update_layout(height=700, showlegend=True, hovermode='x unified')

            st.plotly_chart(fig, use_container_width=True)

            # Event details table
            if rain_events:
                st.subheader("🌧️ Rain Event Details")
                event_df = pd.DataFrame(rain_events).drop(columns=['start_idx', 'end_idx'], errors='ignore')
                st.dataframe(event_df, use_container_width=True)

                # Event summary chart
                fig2 = make_subplots(
                    rows=1, cols=2,
                    subplot_titles=('I&I Volume by Event', 'Peak Flow vs Rainfall')
                )

                fig2.add_trace(
                    go.Bar(x=[f"Event {e['event_num']}" for e in rain_events],
                          y=[e.get('ii_volume_mg', 0) for e in rain_events],
                          name='I&I Volume',
                          marker_color='coral'),
                    row=1, col=1
                )

                fig2.add_trace(
                    go.Scatter(x=[e['total_rain'] for e in rain_events],
                              y=[e.get('peak_flow_gpm', 0) for e in rain_events],
                              mode='markers',
                              marker=dict(size=12, color='steelblue'),
                              name='Peak Flow vs Rain'),
                    row=1, col=2
                )

                fig2.update_xaxes(title_text="Event", row=1, col=1)
                fig2.update_yaxes(title_text="I&I Volume (MG)", row=1, col=1)
                fig2.update_xaxes(title_text="Total Rainfall (in)", row=1, col=2)
                fig2.update_yaxes(title_text="Peak Flow (GPM)", row=1, col=2)
                fig2.update_layout(height=400, showlegend=False)

                st.plotly_chart(fig2, use_container_width=True)

            # Basin comparison if asset data provided
            if asset_file and 'basin_id' in flow_df.columns:
                st.subheader("🏗️ Basin Comparison & Prioritization")

                basin_stats = []
                for basin in flow_df['basin_id'].unique():
                    basin_flow = flow_df[flow_df['basin_id'] == basin]
                    peak_flow = basin_flow['flow_rate_gpm'].max()
                    ii_volume = ((basin_flow['flow_rate_gpm'] - baseline).clip(lower=0).sum() * 60) / 1000000

                    basin_stats.append({
                        'Basin': basin,
                        'Peak Flow (GPM)': peak_flow,
                        'I&I Volume (MG)': ii_volume,
                        'Peak/Baseline Ratio': peak_flow / baseline if baseline > 0 else 0,
                        'Anomaly Count': basin_flow['is_anomaly'].sum() if enable_anomaly else 0
                    })

                basin_df = pd.DataFrame(basin_stats)
                if asset_df is not None:
                    basin_df = basin_df.merge(asset_df, left_on='Basin', right_on='basin_id', how='left')

                    # Calculate priority score
                    basin_df['Priority Score'] = (
                        basin_df['I&I Volume (MG)'] * 50 +
                        basin_df.get('pipe_age_years', 0) * 30 +
                        basin_df.get('pipe_length_miles', 0) * 20
                    )
                    basin_df = basin_df.sort_values('Priority Score', ascending=False)
                    basin_df['Priority Rank'] = range(1, len(basin_df) + 1)

                st.dataframe(basin_df, use_container_width=True)

            # === DOWNLOAD OPTIONS ===
            st.markdown("---")
            st.subheader("💾 Download Results")

            col1, col2, col3 = st.columns(3)

            with col1:
                # Event data CSV
                if rain_events:
                    event_csv = pd.DataFrame(rain_events).to_csv(index=False)
                    st.download_button(
                        "📥 Event Data (CSV)",
                        event_csv,
                        "ii_events.csv",
                        "text/csv"
                    )

            with col2:
                # Complete results JSON
                results_dict = {
                    'analysis_date': datetime.now().isoformat(),
                    'settings': {
                        'rain_threshold': rain_threshold,
                        'baseline_percentile': baseline_percentile,
                        'baseline_flow_gpm': baseline
                    },
                    'events': rain_events,
                    'rdii': rdii_results if rdii_results else {},
                    'summary': {
                        'total_events': len(rain_events),
                        'total_rainfall_in': sum(e['total_rain'] for e in rain_events) if rain_events else 0,
                        'total_ii_volume_mg': sum(e.get('ii_volume_mg', 0) for e in rain_events) if rain_events else 0
                    }
                }
                results_json = json.dumps(results_dict, indent=2, default=str)
                st.download_button(
                    "📥 Full Results (JSON)",
                    results_json,
                    "ii_analysis.json",
                    "application/json"
                )

            with col3:
                # Flow data with calculated fields
                export_df = flow_df.copy()
                export_csv = export_df.to_csv(index=False)
                st.download_button(
                    "📥 Processed Flow Data (CSV)",
                    export_csv,
                    "ii_flow_processed.csv",
                    "text/csv"
                )

            # Second row for client report
            st.markdown("")  # Add spacing
            col4 = st.columns(1)[0]

            with col4:
                # Generate client report
                report_buffer = generate_client_report(
                    rain_events=rain_events,
                    baseline=baseline,
                    flow_df=flow_df,
                    rain_df=rain_df,
                    asset_df=asset_df if asset_df is not None else None,
                    rdii_results=rdii_results,
                    enable_anomaly=enable_anomaly,
                    basin_priority_df=basin_df if (asset_file and 'basin_id' in flow_df.columns) else None
                )

                st.download_button(
                    "📄 Client Report (Word Document)",
                    report_buffer,
                    f"II_Analysis_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Download a professional Word document report ready for client presentation"
                )

else:
    st.info("👆 Upload flow and rainfall data to begin analysis")

    with st.expander("📖 Quick Start Guide"):
        st.markdown("""
        ### How to Use This Tool

        1. **Upload Data**: Provide flow data (required), rainfall data (required), and asset data (optional)
        2. **Configure Settings**: Adjust rain threshold, baseline calculation, and enable advanced features
        3. **Run Analysis**: Click the button to perform comprehensive I&I analysis
        4. **Review Results**: Examine metrics, charts, and event details
        5. **Download**: Export results in CSV or JSON format

        ### Required Data Format

        **Flow Data**: `timestamp, flow_rate_gpm, basin_id (optional)`

        **Rainfall Data**: `timestamp, rainfall_in`

        **Asset Data**: `basin_id, area_acres, pipe_age_years, pipe_length_miles`
        """)
